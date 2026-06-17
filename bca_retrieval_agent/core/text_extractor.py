"""Text extraction from PDF, DOCX, XLSX, and HTML files."""

import logging
from pathlib import Path

from bs4 import BeautifulSoup

from config import EXTRACTED_TEXT_DIR

logger = logging.getLogger(__name__)


def extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts)
    except Exception as exc:
        logger.error("PDF extraction failed for %s: %s", file_path, exc)
        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception as exc2:
            logger.error("pdfplumber fallback failed: %s", exc2)
            return ""


def extract_docx_text(file_path: str) -> str:
    """Extract text from DOCX."""
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)
    except Exception as exc:
        logger.error("DOCX extraction failed for %s: %s", file_path, exc)
        return ""


def extract_xlsx_text(file_path: str) -> str:
    """Extract sheet names, cell values, formulas, and metadata from Excel."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(file_path, data_only=False, read_only=True)
        parts = []

        if wb.properties:
            props = wb.properties
            meta_lines = []
            for attr in ("title", "subject", "creator", "description", "keywords"):
                val = getattr(props, attr, None)
                if val:
                    meta_lines.append(f"{attr}: {val}")
            if meta_lines:
                parts.append("=== Workbook Metadata ===")
                parts.extend(meta_lines)

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"\n=== Sheet: {sheet_name} ===")
            row_count = 0
            for row in ws.iter_rows(max_row=500):
                cells = []
                for cell in row:
                    val = cell.value
                    if val is not None:
                        if isinstance(val, str) and val.startswith("="):
                            cells.append(f"[formula]{val}")
                        else:
                            cells.append(str(val))
                if cells:
                    parts.append(" | ".join(cells))
                    row_count += 1
                    if row_count >= 200:
                        parts.append("... (truncated)")
                        break
        wb.close()
        return "\n".join(parts)
    except Exception as exc:
        logger.error("XLSX extraction failed for %s: %s", file_path, exc)
        return ""


def extract_html_text(file_path: str) -> str:
    """Extract visible text from HTML file."""
    try:
        content = Path(file_path).read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(content, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)
    except Exception as exc:
        logger.error("HTML extraction failed for %s: %s", file_path, exc)
        return ""


def extract_text(file_path: str, file_type: str, run_id: str, document_id: str) -> str:
    """
    Extract text from a document and save to extracted_text directory.
    Returns the extracted text.
    """
    ft = file_type.lower().lstrip(".")
    extractors = {
        "pdf": extract_pdf_text,
        "docx": extract_docx_text,
        "xlsx": extract_xlsx_text,
        "xlsm": extract_xlsx_text,
        "html": extract_html_text,
        "htm": extract_html_text,
    }

    extractor = extractors.get(ft)
    if not extractor:
        logger.warning("No extractor for file type: %s", ft)
        return ""

    text = extractor(file_path)

    out_dir = EXTRACTED_TEXT_DIR / run_id
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{document_id}.txt"
    out_path.write_text(text, encoding="utf-8")

    return text
